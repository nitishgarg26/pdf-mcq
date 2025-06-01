# Final and Reliable Two-Column PDF MCQ Extractor with Enhanced Question Detection

import streamlit as st
import fitz  # PyMuPDF
import pytesseract
from PIL import Image, ImageEnhance
import io
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

st.title("Enhanced 2-Column MCQ OCR Extractor")

def clean_text(text):
    """Clean text by removing null characters and extra whitespace"""
    return text.replace('\x00', '').strip()

def preprocess_ocr_text(ocr_text):
    """
    Enhanced OCR text preprocessing for better question detection
    """
    # Remove null characters and normalize whitespace
    text = re.sub(r'\x00', '', ocr_text)
    text = re.sub(r'\s+', ' ', text)
    
    # Fix common OCR errors in question numbering
    text = re.sub(r'(\d+)[oO](\s)', r'\1.\2', text)  # "1o " -> "1. "
    text = re.sub(r'(\d+)l(\s)', r'\1.\2', text)     # "1l " -> "1. "
    text = re.sub(r'(\d+),(\s)', r'\1.\2', text)     # "1, " -> "1. "
    
    # Fix spacing around question numbers
    text = re.sub(r'(\d+)\.(\w)', r'\1. \2', text)   # "1.What" -> "1. What"
    
    # Remove excessive line breaks but preserve paragraph structure
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text

def enhanced_question_detection(ocr_text):
    """
    Enhanced question detection with multiple pattern matching
    """
    # Multiple regex patterns for different question formats
    patterns = [
        r"(?:^|\n)(\d{1,3}\.)\s*",           # Standard: "1. Question"
        r"(?:^|\n)(Q\d{1,3}\.?)\s*",         # Format: "Q1. Question" or "Q1 Question"
        r"(?:^|\n)(\d{1,3}\))\s*",           # Format: "1) Question"
        r"(?:^|\n)(\(\d{1,3}\))\s*",         # Format: "(1) Question"
        r"(?:^|\n)(Question\s+\d{1,3}\.?)\s*", # Format: "Question 1. Text"
        r"(?:^|\n)(\d{1,3}[-–—]\s*)",        # Format: "1- Question" or "1– Question"
    ]
    
    best_matches = []
    best_count = 0
    
    # Try each pattern and use the one that finds the most questions
    for pattern in patterns:
        matches = re.finditer(pattern, ocr_text, re.MULTILINE | re.IGNORECASE)
        current_matches = list(matches)
        
        if len(current_matches) > best_count:
            best_count = len(current_matches)
            best_matches = current_matches
    
    return best_matches

def split_questions_from_ocr_enhanced(ocr_text):
    """
    Enhanced question splitting with improved pattern detection
    """
    # Clean the text first
    cleaned_text = preprocess_ocr_text(ocr_text)
    
    matches = enhanced_question_detection(cleaned_text)
    
    if not matches:
        # Fallback to simple pattern
        pattern = re.compile(r"(?:^|\n)(\d{1,3}\.)")
        parts = pattern.split(cleaned_text)
        combined = []
        i = 1
        while i < len(parts):
            question_num = parts[i].strip()
            question_text = parts[i + 1].strip() if i + 1 < len(parts) else ""
            combined.append((question_num, question_text))
            i += 2
        return combined
    
    questions = []
    
    for i, match in enumerate(matches):
        question_num = match.group(1).strip()
        start_pos = match.end()
        
        # Find the end position (start of next question or end of text)
        if i + 1 < len(matches):
            end_pos = matches[i + 1].start()
        else:
            end_pos = len(cleaned_text)
        
        question_text = cleaned_text[start_pos:end_pos].strip()
        
        # Remove answer choices if they follow the pattern (A), (B), etc.
        question_text = re.sub(r'\s*\([A-E]\)[^(]*(?=\([A-E]\)|\s*$)', '', question_text)
        
        questions.append((question_num, question_text))
    
    return questions

def detect_question_boundaries_advanced(img, data):
    """
    Advanced question boundary detection using OCR confidence and positioning
    """
    question_positions = []
    
    # Find question number patterns with their positions
    for i, text in enumerate(data['text']):
        text_clean = text.strip()
        confidence = int(data['conf'][i]) if data['conf'][i] != '-1' else 0
        
        # Multiple patterns for question detection
        patterns = [
            r'^\d{1,3}\.$',
            r'^Q\d{1,3}\.?$',
            r'^\d{1,3}\)$',
            r'^\(\d{1,3}\)$'
        ]
        
        for pattern in patterns:
            if re.match(pattern, text_clean) and confidence > 30:  # Confidence threshold
                top = data['top'][i]
                left = data['left'][i]
                width = data['width'][i]
                height = data['height'][i]
                
                # Avoid duplicate detections (same position)
                is_duplicate = any(abs(pos['top'] - top) < 10 and abs(pos['left'] - left) < 20 
                                 for pos in question_positions)
                
                if not is_duplicate:
                    question_positions.append({
                        'top': top,
                        'left': left,
                        'width': width,
                        'height': height,
                        'text': text_clean,
                        'confidence': confidence
                    })
                break
    
    # Sort by vertical position (top to bottom)
    question_positions.sort(key=lambda x: (x['top'], x['left']))
    
    return question_positions

def trim_horizontal(img):
    """Trim empty left/right margins"""
    gray = img.convert("L")
    bbox = gray.point(lambda x: 0 if x > 240 else 255, '1').getbbox()
    if bbox:
        return img.crop((bbox[0], 0, bbox[2], img.height))
    return img

def trim_question_number_horizontal(img, data, question_top_global):
    """
    Trim the question number horizontally (remove from left side)
    """
    question_number_right = 0
    
    # Find question number position within the image
    for i, text in enumerate(data['text']):
        text_clean = text.strip()
        top = data['top'][i]
        left = data['left'][i]
        width = data['width'][i]
        
        # Check if this is a question number near the expected position
        patterns = [r'^\d{1,3}\.$', r'^Q\d{1,3}\.?$', r'^\d{1,3}\)$', r'^\(\d{1,3}\)$']
        
        for pattern in patterns:
            if re.match(pattern, text_clean) and abs(top - question_top_global) < 30:
                # Find the rightmost edge of the question number + some padding
                question_number_right = max(question_number_right, left + width + 10)
                break
    
    # If we found a question number, crop horizontally from its right edge
    if question_number_right > 0 and question_number_right < img.width - 50:
        img = img.crop((question_number_right, 0, img.width, img.height))
    
    return img

def set_landscape_and_borders(doc):
    """Set document to landscape and add borders to all tables"""
    # Set document to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Set margins for better use of space
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add borders to all tables
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            borders.append(border)
        tblPr.append(borders)

def validate_ocr_quality(data):
    """Check OCR quality based on confidence scores"""
    confidences = [int(conf) for conf in data['conf'] if conf != '-1' and int(conf) > 0]
    if not confidences:
        return False
    avg_confidence = sum(confidences) / len(confidences)
    return avg_confidence > 50  # Threshold for acceptable quality

@st.cache_data
def extract_questions_from_columns_enhanced(pdf_bytes):
    """
    Enhanced question extraction with improved detection and horizontal question number trimming
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted = []
    processing_stats = {"total_pages": len(doc), "questions_found": 0, "low_quality_pages": 0}

    for page_num, page in enumerate(doc):
        width = page.rect.width
        height = page.rect.height
        mid_x = width / 2

        columns = [
            fitz.Rect(0, 0, mid_x, height),
            fitz.Rect(mid_x, 0, width, height)
        ]

        for col_num, col in enumerate(columns):
            try:
                pix = page.get_pixmap(clip=col, dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")

                # Enhanced image preprocessing
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(1.5)
                
                # Sharpen the image for better OCR
                enhancer = ImageEnhance.Sharpness(img)
                img = enhancer.enhance(1.2)

                # Get OCR text and data
                ocr_text = pytesseract.image_to_string(img)
                data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
                
                # Check OCR quality
                if not validate_ocr_quality(data):
                    processing_stats["low_quality_pages"] += 1
                    st.warning(f"Low OCR quality detected on page {page_num + 1}, column {col_num + 1}")
                
                # Enhanced question detection
                questions = split_questions_from_ocr_enhanced(ocr_text)
                question_positions = detect_question_boundaries_advanced(img, data)

                # Use the better of the two detection methods
                if len(question_positions) > len(questions):
                    question_tops = [pos['top'] for pos in question_positions]
                else:
                    # Fallback to original method if advanced detection finds fewer questions
                    seen = set()
                    question_tops = []
                    for i, text in enumerate(data['text']):
                        text_clean = text.strip()
                        if re.match(r'^\d{1,3}\.$', text_clean):
                            top = data['top'][i]
                            if top not in seen:
                                question_tops.append(top)
                                seen.add(top)
                    question_tops = sorted(set(question_tops))

                question_tops.append(img.height)

                # Extract and crop questions
                for i in range(len(question_tops) - 1):
                    y1 = max(0, question_tops[i] - 15)  # More generous spacing
                    y2 = question_tops[i + 1]
                    
                    if y2 - y1 < 20:  # Skip very small regions
                        continue
                        
                    cropped_img = img.crop((0, y1, img.width, y2))
                    
                    # Trim question number horizontally (from left side)
                    cropped_img = trim_question_number_horizontal(cropped_img, data, question_tops[i])
                    
                    cropped_img = trim_horizontal(cropped_img)
                    
                    # Skip if cropped image is too small
                    if cropped_img.width < 50 or cropped_img.height < 20:
                        continue
                    
                    buf = io.BytesIO()
                    cropped_img.save(buf, format="PNG")

                    q_text = "(image only)"
                    confidence_score = 0
                    
                    if i < len(questions):
                        # Remove question number from text as well
                        question_text = questions[i][1].strip()
                        q_text = question_text
                    
                    if i < len(question_positions):
                        confidence_score = question_positions[i]['confidence']

                    extracted.append({
                        "text": clean_text(q_text),
                        "image": buf.getvalue(),
                        "confidence": confidence_score,
                        "page": page_num + 1,
                        "column": col_num + 1
                    })
                    processing_stats["questions_found"] += 1

            except Exception as e:
                st.error(f"Error processing page {page_num + 1}, column {col_num + 1}: {str(e)}")
                continue

    return extracted, processing_stats

def generate_word_enhanced(questions):
    """Generate Word document with landscape layout and bordered table"""
    doc = Document()
    
    # Create table directly without any headers or titles - 5 columns
    table = doc.add_table(rows=0, cols=5)  # No header row
    
    for idx, q in enumerate(questions, 1):
        row = table.add_row().cells
        
        # Add image (question without number)
        para = row[0].paragraphs[0]
        run = para.add_run()
        try:
            run.add_picture(io.BytesIO(q["image"]), width=Inches(3.0))
        except Exception as e:
            para.add_run(f"[Image Error: {str(e)}]")
        
        # Add question text (without question number)
        row[1].text = q['text']
        
        # Empty cells for options (to be filled manually)
        for i in range(2, 5):
            row[i].text = ""
    
    # Set landscape orientation and add borders
    set_landscape_and_borders(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.markdown("""
### Enhanced MCQ Extractor Features:
- **Multiple Question Formats**: Supports various numbering styles (1., Q1., 1), (1), etc.)
- **Improved OCR**: Enhanced preprocessing and error correction
- **Quality Validation**: OCR confidence scoring and quality checks
- **Advanced Detection**: Multiple pattern matching algorithms
- **Horizontal Question Number Trimming**: Removes question numbers from left side of images
- **Landscape Layout**: Word document in landscape orientation with bordered tables
""")

# File upload with size validation
pdf_file = st.file_uploader("Upload a 2-column MCQ PDF", type=["pdf"])

if pdf_file:
    # Validate file size (50MB limit)
    if pdf_file.size > 50 * 1024 * 1024:
        st.error("File too large. Please upload a PDF smaller than 50MB.")
    else:
        try:
            pdf_bytes = pdf_file.read()
            
            # Store extraction date
            from datetime import datetime
            st.session_state["extraction_date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            with st.spinner("Processing PDF with enhanced OCR question detection..."):
                questions, stats = extract_questions_from_columns_enhanced(pdf_bytes)

            if not questions:
                st.warning("No questions detected. Please check if the PDF contains numbered questions in a two-column format.")
            else:
                # Display processing statistics
                st.success(f"✅ Processing Complete!")
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Questions Found", stats["questions_found"])
                with col2:
                    st.metric("Pages Processed", stats["total_pages"])
                with col3:
                    st.metric("Low Quality Pages", stats["low_quality_pages"])
                
                # Display questions with enhanced information
                st.subheader("Extracted Questions Preview")
                
                # Add filters
                show_confidence = st.checkbox("Show confidence scores", value=True)
                min_confidence = st.slider("Minimum confidence threshold", 0, 100, 0)
                
                # Filter questions by confidence
                filtered_questions = [q for q in questions if q.get('confidence', 0) >= min_confidence]
                
                if not filtered_questions:
                    st.warning(f"No questions meet the confidence threshold of {min_confidence}%")
                else:
                    for i, q in enumerate(filtered_questions, 1):
                        with st.expander(f"Question {i} (Page {q.get('page', 'N/A')}, Column {q.get('column', 'N/A')})"):
                            col1, col2 = st.columns([2, 1])
                            
                            with col1:
                                st.image(q["image"], caption=f"Question {i} (Number Trimmed Horizontally)", use_container_width=True)
                            
                            with col2:
                                st.markdown(f"**Text:** {q['text']}")
                                if show_confidence and q.get('confidence', 0) > 0:
                                    confidence_color = "🟢" if q['confidence'] > 70 else "🟡" if q['confidence'] > 40 else "🔴"
                                    st.markdown(f"**Confidence:** {confidence_color} {q['confidence']}%")

                    # Generate and offer download
                    st.subheader("Download Results")
                    
                    with st.spinner("Generating Word document with landscape layout and borders..."):
                        docx_file = generate_word_enhanced(filtered_questions)
                    
                    st.download_button(
                        label=f"📄 Download Word Document ({len(filtered_questions)} questions)",
                        data=docx_file,
                        file_name=f"extracted_mcqs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    # Additional export options
                    if st.button("📊 Export Processing Statistics"):
                        stats_text = f"""
Processing Statistics:
- Total Pages: {stats['total_pages']}
- Questions Found: {stats['questions_found']}
- Low Quality Pages: {stats['low_quality_pages']}
- Average Questions per Page: {stats['questions_found'] / stats['total_pages']:.1f}
- Extraction Date: {st.session_state.get('extraction_date', 'N/A')}
                        """
                        st.download_button(
                            label="Download Statistics",
                            data=stats_text,
                            file_name=f"extraction_stats_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain"
                        )

        except Exception as e:
            st.error(f"❌ Processing error: {str(e)}")
            st.info("Please ensure the PDF is not corrupted and contains readable text.")

# Footer with usage tips
st.markdown("""
---
### 💡 Tips for Best Results:
- Use high-quality scanned PDFs with clear text
- Ensure questions are numbered consistently (1., 2., 3., etc.)
- Two-column layout works best
- Avoid heavily compressed or low-resolution files
- Check confidence scores to identify potential OCR errors
- Question numbers are automatically trimmed horizontally from the left side
- **New**: Word document now in landscape format with bordered tables for better readability
""")
